from pathlib import Path
import io
from concurrent.futures import ThreadPoolExecutor, as_completed

from langchain_core.documents import Document
from pypdf import PdfReader
import docx as docx_lib


def _ocr_page(data: bytes, page_number: int):
    try:
        import pytesseract
        from pdf2image import convert_from_bytes

        images = convert_from_bytes(
            data, dpi=150, first_page=page_number, last_page=page_number,
            fmt="jpeg", thread_count=1
        )
        if not images:
            return page_number, ""

        text = pytesseract.image_to_string(
            images[0], config="--psm 6", timeout=60
        ).strip()
        return page_number, text
    except Exception as e:
        print(f"[OCR] Page {page_number} failed: {e}")
        return page_number, ""


def load_pdf(file):
    data = file.read() if hasattr(file, "read") else Path(file).read_bytes()
    reader = PdfReader(io.BytesIO(data))
    documents = []
    pages_needing_ocr = []

    # Fast path: normal PDF text extraction first.
    for page_number, page in enumerate(reader.pages, start=1):
        try:
            text = (page.extract_text() or "").strip()
        except Exception as e:
            print(f"[PDF] Extraction failed on page {page_number}: {e}")
            text = ""

        if text:
            documents.append(Document(
                page_content=text,
                metadata={
                    "source": getattr(file, "name", str(file)),
                    "file_type": "pdf",
                    "page": page_number,
                    "ocr": False,
                },
            ))
        else:
            pages_needing_ocr.append(page_number)

    # OCR only pages that have no text layer.
    if pages_needing_ocr:
        print(
            f"[PDF] OCR required for {len(pages_needing_ocr)} "
            f"of {len(reader.pages)} pages."
        )

        results = {}
        workers = min(4, len(pages_needing_ocr))

        with ThreadPoolExecutor(max_workers=workers) as executor:
            futures = {
                executor.submit(_ocr_page, data, page): page
                for page in pages_needing_ocr
            }
            for future in as_completed(futures):
                page, text = future.result()
                if text:
                    results[page] = text

        for page_number in pages_needing_ocr:
            text = results.get(page_number, "")
            if text:
                documents.append(Document(
                    page_content=text,
                    metadata={
                        "source": getattr(file, "name", str(file)),
                        "file_type": "pdf",
                        "page": page_number,
                        "ocr": True,
                    },
                ))

    documents.sort(key=lambda d: d.metadata.get("page", 0))
    print(f"[PDF] Loaded {len(documents)}/{len(reader.pages)} pages.")
    return documents


def load_docx(file):
    doc = docx_lib.Document(file)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    if not text:
        return []
    return [Document(
        page_content=text,
        metadata={"source": getattr(file, "name", str(file)), "file_type": "docx"}
    )]


def load_txt(file):
    raw = file.read()
    text = raw if isinstance(raw, str) else raw.decode("utf-8", errors="ignore")
    text = text.strip()
    if not text:
        return []
    return [Document(
        page_content=text,
        metadata={"source": getattr(file, "name", str(file)), "file_type": "txt"}
    )]


def extract_documents(file):
    suffix = Path(getattr(file, "name", str(file))).suffix.lower()
    if suffix == ".pdf":
        return load_pdf(file)
    if suffix == ".docx":
        return load_docx(file)
    if suffix == ".txt":
        return load_txt(file)
    raise ValueError(f"Unsupported file type: {suffix}")


def extract_text(file, suffix: str):
    return "\n\n".join(d.page_content for d in extract_documents(file))


def load_documents(directory: str):
    documents = []
    for path in Path(directory).glob("*"):
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            documents.extend(load_pdf(path))
        elif suffix == ".docx":
            documents.extend(load_docx(path))
        elif suffix == ".txt":
            documents.extend(load_txt(path))
    return documents